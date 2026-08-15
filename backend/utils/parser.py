import os
import re

import pdfplumber
from docx import Document

from data.skills import SKILLS


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted resume text while
    preserving useful line breaks.
    """

    if not text:
        return ""

    # Remove null bytes
    text = text.replace("\x00", " ")

    # Remove invalid control characters
    text = re.sub(
        r"[\x01-\x08\x0B\x0C\x0E-\x1F\x7F]",
        " ",
        text
    )

    # Normalize line endings
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Clean each line
    lines = []

    for line in text.split("\n"):

        line = re.sub(
            r"[ \t]+",
            " ",
            line
        )

        line = line.strip()

        if line:
            lines.append(line)

    # Remove excessive blank lines
    cleaned = "\n".join(lines)

    return cleaned.strip()


# ============================================================
# PDF EXTRACTION
# ============================================================

def extract_pdf_text(pdf_path: str) -> str:
    """
    Extract text from a PDF resume using pdfplumber.
    """

    if not os.path.exists(pdf_path):

        raise FileNotFoundError(
            f"PDF file not found: {pdf_path}"
        )

    text_parts = []

    try:

        with pdfplumber.open(pdf_path) as pdf:

            for page in pdf.pages:

                extracted = page.extract_text()

                if extracted:
                    text_parts.append(
                        extracted
                    )

    except Exception as e:

        raise ValueError(
            f"Failed to extract PDF text: {str(e)}"
        )

    return "\n".join(text_parts)


# ============================================================
# DOCX EXTRACTION
# ============================================================

def extract_docx_text(docx_path: str) -> str:
    """
    Extract text from a DOCX resume.

    Extracts:
    - Paragraphs
    - Tables
    """

    if not os.path.exists(docx_path):

        raise FileNotFoundError(
            f"DOCX file not found: {docx_path}"
        )

    text_parts = []

    try:

        document = Document(
            docx_path
        )

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                text_parts.append(
                    text
                )

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:

                        row_text.append(
                            cell_text
                        )

                if row_text:

                    text_parts.append(
                        " | ".join(row_text)
                    )

    except Exception as e:

        raise ValueError(
            f"Failed to extract DOCX text: {str(e)}"
        )

    return "\n".join(text_parts)


# ============================================================
# MAIN TEXT EXTRACTION
# ============================================================

def extract_text(file_path: str) -> str:
    """
    Automatically extract text from PDF or DOCX.
    """

    if not os.path.exists(file_path):

        raise FileNotFoundError(
            f"File not found: {file_path}"
        )

    extension = os.path.splitext(
        file_path
    )[1].lower()

    if extension == ".pdf":

        text = extract_pdf_text(
            file_path
        )

    elif extension == ".docx":

        text = extract_docx_text(
            file_path
        )

    else:

        raise ValueError(
            "Unsupported file type. "
            "Only PDF and DOCX files are supported."
        )

    text = clean_text(
        text
    )

    if not text:

        raise ValueError(
            "No readable text found in the resume."
        )

    return text


# ============================================================
# SKILL EXTRACTION
# ============================================================

def extract_skill(text: str) -> list:
    """
    Extract technical skills defined in
    data/skills.py.
    """

    if not text:
        return []

    normalized_text = text.lower()

    found_skills = []

    for skill in SKILLS:

        skill_lower = skill.lower()

        # Prevent partial matches.
        # Example:
        # Java should not match JavaScript.
        pattern = (
            r"(?<!\w)"
            + re.escape(skill_lower)
            + r"(?!\w)"
        )

        if re.search(
            pattern,
            normalized_text
        ):

            found_skills.append(
                skill
            )

    return found_skills